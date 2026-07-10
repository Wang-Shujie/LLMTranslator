"""文档翻译：逐段并发翻译 .docx/.txt，保留段落结构。

TranslationGranularity 策略接口（PerParagraph 实现，WholeDocument 预留 B）。
DocxHandler 用 python-docx 逐段抽取+写回；TxtHandler 用标准库空行分块。
"""
from __future__ import annotations

import asyncio
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import Callable


class Segment:
    """一段文本 + 写回回调（翻译后调用 set_translation 写回原文档对象）。"""

    def __init__(self, text: str, setter: Callable[[str], None]) -> None:
        self.text = text
        self._setter = setter

    def set_translation(self, translation: str) -> None:
        self._setter(translation)


# ---- 翻译粒度策略（预留 B 的扩展点）----

class TranslationGranularity(ABC):
    """segments → 译文列表，1:1 对应。"""

    @abstractmethod
    async def translate(self, segments: list[str], provider, src: str, tgt: str) -> list[str]:
        ...


class PerParagraphGranularity(TranslationGranularity):
    """逐段并发翻译（v1，方案①）。"""

    def __init__(self, concurrency: int = 8) -> None:
        self._sem = asyncio.Semaphore(concurrency)

    async def translate(self, segments: list[str], provider, src: str, tgt: str) -> list[str]:
        async def one(text: str) -> str:
            async with self._sem:
                parts: list[str] = []
                async for tok in provider.translate(text, src, tgt):
                    parts.append(tok)
                return "".join(parts)
        return await asyncio.gather(*[one(s) for s in segments])


class WholeDocumentGranularity(TranslationGranularity):
    """整篇翻译（预留，方案② B），v1 不实现。"""

    async def translate(self, segments: list[str], provider, src: str, tgt: str) -> list[str]:
        raise NotImplementedError("整篇翻译：后续版本实现")


# ---- 格式处理器 ----

class TxtHandler:
    """纯文本：按空行分块。"""

    def __init__(self) -> None:
        self._blocks: list[str] = []

    def extract(self, path: str) -> list[Segment]:
        text = Path(path).read_text(encoding="utf-8")
        self._blocks = text.split("\n\n")
        return [
            Segment(b, lambda t, i=i: self._blocks.__setitem__(i, t))
            for i, b in enumerate(self._blocks)
        ]

    def save(self, out_path: str) -> None:
        Path(out_path).write_text("\n\n".join(self._blocks), encoding="utf-8")


class DocxHandler:
    """Word .docx：python-docx 逐段抽取（正文+表格）+ 写回。段内 run 格式为已知边界。"""

    def __init__(self) -> None:
        self._doc = None

    def extract(self, path: str) -> list[Segment]:
        from docx import Document as DocxDocument
        self._doc = DocxDocument(path)
        segments: list[Segment] = []
        for para in self._doc.paragraphs:
            if para.text.strip():
                segments.append(Segment(para.text, self._make_setter(para)))
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            segments.append(Segment(para.text, self._make_setter(para)))
        return segments

    @staticmethod
    def _make_setter(para):
        def setter(t: str) -> None:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = t
            else:
                para.add_run(t)
        return setter

    def save(self, out_path: str) -> None:
        self._doc.save(out_path)


def handler_for_format(path: str):
    """根据文件扩展名返回对应的 Handler。"""
    ext = Path(path).suffix.lower()
    if ext == ".txt":
        return TxtHandler()
    if ext == ".docx":
        return DocxHandler()
    raise ValueError(f"不支持的格式：{ext}（v1 仅支持 .docx/.txt）")


# ---- 编排 ----

class DocumentTranslator:
    """文档翻译编排：extract → translate → write-back → save。"""

    def __init__(self, granularity: TranslationGranularity) -> None:
        self._granularity = granularity

    async def translate_document(
        self, path: str, out_path: str, provider, src: str, tgt: str
    ) -> int:
        """翻译文档，返回段数。失败抛异常。"""
        handler = handler_for_format(path)
        segments = handler.extract(path)
        if not segments:
            raise ValueError("文档为空或无文字")
        texts = [s.text for s in segments]
        translations = await self._granularity.translate(texts, provider, src, tgt)
        for seg, tr in zip(segments, translations):
            seg.set_translation(tr)
        handler.save(out_path)
        return len(segments)
