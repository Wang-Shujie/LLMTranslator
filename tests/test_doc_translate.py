import asyncio
import pytest
from llm_translator.core.doc_translate import (
    Segment, TranslationGranularity, PerParagraphGranularity,
    WholeDocumentGranularity, TxtHandler, DocxHandler, DocumentTranslator,
    handler_for_format,
)


class _FakeProvider:
    """Fake provider that translates by appending '_T'."""
    async def login(self): pass
    async def translate(self, text, src, tgt):
        for tok in [text + "_T"]:
            yield tok


# ---- Granularity ----

@pytest.mark.asyncio
async def test_per_paragraph_translates_each_segment():
    g = PerParagraphGranularity(concurrency=4)
    result = await g.translate(["hello", "world"], _FakeProvider(), "en", "zh")
    assert result == ["hello_T", "world_T"]


def test_whole_document_raises_not_implemented():
    g = WholeDocumentGranularity()
    loop = asyncio.new_event_loop()
    try:
        with pytest.raises(NotImplementedError):
            loop.run_until_complete(
                asyncio.wait_for(g.translate(["x"], _FakeProvider(), "en", "zh"), 1)
            )
    finally:
        loop.close()


# ---- TxtHandler ----

def test_txt_handler_roundtrip(tmp_path):
    f = tmp_path / "test.txt"
    f.write_text("Hello\n\nWorld", encoding="utf-8")
    handler = TxtHandler()
    segments = handler.extract(str(f))
    assert len(segments) == 2
    assert segments[0].text == "Hello"
    segments[0].set_translation("你好")
    segments[1].set_translation("世界")
    out = tmp_path / "test_zh.txt"
    handler.save(str(out))
    assert out.read_text(encoding="utf-8") == "你好\n\n世界"


def test_txt_handler_empty_blocks_preserved(tmp_path):
    f = tmp_path / "test.txt"
    f.write_text("A\n\n\n\nB", encoding="utf-8")
    handler = TxtHandler()
    segments = handler.extract(str(f))
    segments[0].set_translation("X")
    segments[-1].set_translation("Y")
    out = tmp_path / "out.txt"
    handler.save(str(out))
    content = out.read_text(encoding="utf-8")
    assert "X" in content and "Y" in content


# ---- DocxHandler ----

def test_docx_handler_roundtrip(tmp_path):
    from docx import Document
    f = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("Goodbye")
    doc.save(str(f))

    handler = DocxHandler()
    segments = handler.extract(str(f))
    assert len(segments) == 2
    assert segments[0].text == "Hello world"
    segments[0].set_translation("你好世界")
    segments[1].set_translation("再见")
    out = tmp_path / "test_zh.docx"
    handler.save(str(out))

    doc2 = Document(str(out))
    texts = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert texts == ["你好世界", "再见"]


def test_docx_handler_table_cells(tmp_path):
    from docx import Document
    f = tmp_path / "test.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    doc.save(str(f))

    handler = DocxHandler()
    segments = handler.extract(str(f))
    texts = [s.text for s in segments]
    assert "Name" in texts and "Age" in texts
    for s in segments:
        s.set_translation(s.text + "_T")
    out = tmp_path / "out.docx"
    handler.save(str(out))
    doc2 = Document(str(out))
    assert doc2.tables[0].cell(0, 0).text.strip().endswith("_T")


# ---- handler_for_format ----

def test_handler_for_format_txt():
    assert isinstance(handler_for_format("file.txt"), TxtHandler)


def test_handler_for_format_docx():
    assert isinstance(handler_for_format("file.docx"), DocxHandler)


def test_handler_for_format_unsupported():
    with pytest.raises(ValueError):
        handler_for_format("file.pdf")
